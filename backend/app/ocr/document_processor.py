import io
import os
import logging
import tempfile
import pymupdf
import pdfplumber
import openpyxl
import docx
from PIL import Image
from app.ai.gemini_client import analyze_image, analyze_text
from app.ai.groq_client import analyze_document

logger = logging.getLogger("civilai.document_processor")


# ── PDF text extraction ────────────────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text


# ── Camelot PDF table extraction ───────────────────────────────────────────────

def extract_pdf_tables(file_bytes: bytes) -> str:
    """
    Extract tables from a PDF using camelot-py (stream mode — no Ghostscript needed).
    Returns tables serialised as pipe-delimited text, separated by blank lines.
    Falls back silently if camelot is unavailable or the PDF has no detectable tables.
    """
    try:
        import camelot
    except ImportError:
        logger.debug("camelot not installed — skipping table extraction")
        return ""

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        # Try stream mode first (works without Ghostscript)
        tables = camelot.read_pdf(tmp_path, flavor="stream", pages="all", suppress_stdout=True)

        # If stream found nothing useful, try lattice (requires Ghostscript — may fail)
        if not tables or all(t.df.empty for t in tables):
            try:
                tables = camelot.read_pdf(tmp_path, flavor="lattice", pages="all", suppress_stdout=True)
            except Exception:
                pass

        if not tables:
            return ""

        blocks: list[str] = []
        for i, table in enumerate(tables):
            if table.df.empty:
                continue
            accuracy = getattr(table, "accuracy", 0)
            if accuracy < 30:   # skip very low-confidence extractions
                continue
            # Convert DataFrame to pipe-delimited text
            rows = table.df.values.tolist()
            lines = [" | ".join(str(cell).strip() for cell in row) for row in rows if any(str(c).strip() for c in row)]
            if lines:
                blocks.append(f"[Table {i + 1}]\n" + "\n".join(lines))

        result = "\n\n".join(blocks)
        logger.info("camelot extracted %d table(s), %d chars", len(blocks), len(result))
        return result

    except Exception as exc:
        logger.debug("camelot table extraction failed: %s", exc)
        return ""
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


# ── Other format extractors ────────────────────────────────────────────────────

def extract_pdf_images(file_bytes: bytes) -> list:
    images = []
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    for page in doc:
        pix = page.get_pixmap()
        images.append(pix.tobytes("png"))
    return images


def extract_excel(file_bytes: bytes) -> str:
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    except Exception:
        # openpyxl only reads the OOXML .xlsx format — a real legacy .xls
        # (binary/BIFF) file raises here, so fall back to xlrd for that case.
        return _extract_legacy_xls(file_bytes)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        text += f"Sheet: {sheet}\n"
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                text += " | ".join(cells) + "\n"
    return text


def _extract_legacy_xls(file_bytes: bytes) -> str:
    try:
        import xlrd
    except ImportError:
        logger.warning("xlrd not installed — cannot read legacy .xls files")
        return ""
    wb = xlrd.open_workbook(file_contents=file_bytes)
    text = ""
    for sheet in wb.sheets():
        text += f"Sheet: {sheet.name}\n"
        for row_idx in range(sheet.nrows):
            cells = [str(c) for c in sheet.row_values(row_idx) if c not in (None, "")]
            if cells:
                text += " | ".join(cells) + "\n"
    return text


def extract_word(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception:
        # python-docx only reads the OOXML .docx format — a real legacy .doc
        # (binary) file raises here; no lightweight pure-Python reader exists
        # for it, so surface empty text and let the caller's "could not
        # extract text" message tell the user to re-save as .docx.
        logger.warning("Could not open .doc/.docx file — likely a legacy binary .doc")
        return ""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # A worker/team roster in a Word doc is almost always a table, not prose
    # paragraphs — doc.paragraphs alone misses it entirely.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ── Main processor ─────────────────────────────────────────────────────────────

def process_document(file_bytes: bytes, filename: str, prompt: str = None) -> dict:
    ext  = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    text = ""

    try:
        if ext == "pdf":
            # Primary: pdfplumber for text
            text = extract_pdf_text(file_bytes)

            # Enhancement: camelot table extraction — appended to pdfplumber text
            table_text = extract_pdf_tables(file_bytes)
            if table_text:
                separator = "\n\n--- EXTRACTED TABLES ---\n\n"
                text = (text or "") + separator + table_text

            # Fallback: Gemini vision OCR for scanned (image-only) PDFs
            if not text.strip():
                images = extract_pdf_images(file_bytes)
                if images:
                    text = analyze_image(images[0], "Extract all text from this document")

        elif ext in ("xlsx", "xls"):
            text = extract_excel(file_bytes)

        elif ext in ("docx", "doc"):
            text = extract_word(file_bytes)

        elif ext in ("png", "jpg", "jpeg", "webp"):
            text = analyze_image(file_bytes, "Extract all text and data from this image")

        elif ext == "csv":
            import csv as _csv
            decoded = file_bytes.decode("utf-8-sig", errors="replace")
            reader  = _csv.reader(io.StringIO(decoded))
            text    = "\n".join(" | ".join(row) for row in reader)
    except Exception as exc:
        # A corrupted file or a mismatched/renamed extension (e.g. a .xls
        # saved with a .xlsx name) would otherwise crash the whole request
        # with a raw 500. Degrade to empty text so the caller's own
        # "could not extract text" handling takes over instead.
        logger.warning("Extraction failed for %s (.%s): %s", filename, ext, exc)
        text = ""

    analysis = None
    if prompt and text:
        analysis = analyze_document(text, prompt)

    return {
        "filename":       filename,
        "extracted_text": text,
        "analysis":       analysis,
    }
